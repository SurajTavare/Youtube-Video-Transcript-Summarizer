from flask import Flask, request ,send_file ,jsonify
from youtube_transcript_api import YouTubeTranscriptApi, CouldNotRetrieveTranscript, TranscriptsDisabled, VideoUnavailable
from transformers import pipeline

# **********************-----------TRANSLATION import Modules------------**************************************************
# import statement for root window
import tkinter as tk
from tkinter import ttk
# import statement for translation
from googletrans import Translator, LANGUAGES

# **********************-----------DOWNLOAD import Modules------------**************************************************
import yt_dlp
from threading import Lock
import os

# **********************-----------FILE SAVE import Modules------------**************************************************
from pytube import YouTube
from reportlab.pdfgen import canvas 
from reportlab.pdfbase.ttfonts import TTFont 
from reportlab.pdfbase import pdfmetrics 
from reportlab.lib import colors 
from reportlab.lib.pagesizes import letter
from langdetect import detect
from docx import Document

#************************----------------NOTES import---------------------------------************************************************
import google.generativeai as genai



import re
import requests


app = Flask(__name__)

progress = 0
progress_lock = Lock()
translated_text = None
notes= None
@app.get('/summary')
def summary_api():
    global text
    global video_id
    try:
        url = request.args.get('url', '')
        if not url:
            return jsonify({"error": "URL is required"}), 400
        print(type(url))
        print("url===", url)
        video_id = url.split('=')[1]
        try:
            transcript = get_transcript(video_id)
        except (CouldNotRetrieveTranscript, TranscriptsDisabled, VideoUnavailable) as e:
            return "Could not retrieve transcript: could not retrive summary for this video", 400
        
        summary = get_summary(transcript)
        text = summary
        return  summary, 200

    except Exception as e:
        return  "Could not retrieve transcript: could not retrive summary for this video ", 500



def get_transcript(video_id):
    transcript_list = YouTubeTranscriptApi.get_transcript(video_id)
    transcript = ' '.join([d['text'] for d in transcript_list])
    return transcript

def get_summary(transcript):
    summariser = pipeline('summarization')
    summary = ''
    for i in range(0, (len(transcript)//1000)+1):
        summary_text = summariser(transcript[i*1000:(i+1)*1000])[0]['summary_text']
        summary = summary + summary_text + ' '
    return summary
    
# **************************************TRANSLATION FUNCTION*****************************************
@app.get('/translate')
def translate_text():
    return abc(text)

def abc(txt):
    # root window code
        root = tk.Tk()
        root.title("Language Translator")
        root.geometry('500x300')  # Set the window size
        root.configure(bg="black")

        style = ttk.Style()
        style.theme_use('clam')

        selected_language = tk.StringVar(root)
        selected_language.set("English")  # default value

        # Language list
        languages = ['Marathi', 'Hindi', 'Bengali', 'Telugu', 'Tamil', 'Gujarati' , 'Spanish', 'French', 'German', 'Chinese']

        language_dropdown = ttk.Combobox(root, textvariable=selected_language, values=languages, state="readonly")
        language_dropdown.pack(pady=20, padx=10)

        label = tk.Label(root, text="Select a language to translate:", font=('times new roman', 15))
        label.config(bg="violet")
        label.pack(pady=(10, 0))

        translate_button = tk.Button(root, text="Translate", command=lambda: [ root.destroy() ])
        translate_button.pack(pady=10, padx=10)
        translate_button.config(bg='aquamarine', fg='black', font=('times new roman', 15, ))
        root.mainloop()
      
# *****************************translation logic*******************************8
        try:
        # Translation logic
            global translated_text
            translator = Translator()
            if len(txt)>4000:
                midpoint = len(txt) // 2
                split_point = txt.rfind(' ', 0, midpoint)
    
                if split_point == -1:  # No space found, just split in the middle
                    split_point = midpoint
    
    # ********----------Split the text into two parts-------****************************
                part1 = txt[:split_point].strip()
                part2 = txt[split_point:].strip()
                translation1 = translator.translate(part1, dest=selected_language.get())
                translation2 = translator.translate(part2, dest=selected_language.get())
                # global translated_text
                translated_text = "................................................" + selected_language.get() + "..............................................................."  + "\n\n\n\n" + translation1.text + translation2.text
                return translated_text
            else:
                translation = translator.translate(txt, dest=selected_language.get())
                
                translated_text = "................................................." + selected_language.get() + "..............................................................." + "\n\n\n\n\n" +  translation.text
                return translated_text
        
        except Exception as e:
            return f"Error during translation: {str(e)}", 500


# ********************************************DOWNLOAD FUNCTION***************************************************************
@app.get('/download')
def download():
    url = request.args.get('url', '')
    if not url:
        return "No URL provided", 400

    global progress
    progress = 0  # Reseting progress 

    ydl_opts = {
        'format': 'bestvideo+bestaudio[ext=m4a]/best',
        'ffmpeg_location': "C:/Users/suraj/Downloads/ffmpeg-2024-08-26-git-98610fe95f-essentials_build/ffmpeg-2024-08-26-git-98610fe95f-essentials_build/bin/ffmpeg.exe",
        'merge_output_format': 'mp4',
        'outtmpl': 'downloads/%(title)s.%(ext)s',
        'progress_hooks': [progress_hook]  # Add progress hook
    }

    with yt_dlp.YoutubeDL(ydl_opts) as ydl:
        info = ydl.extract_info(url, download=False)
        filename = ydl.prepare_filename(info)
    file_path=filename
    if os.path.exists(file_path):
        print("File already exists.")
        print("filename",filename)
        return jsonify({"filename":filename,
                        "text":"video already exists"}),800
    
    with yt_dlp.YoutubeDL(ydl_opts) as ydl:
        info = ydl.extract_info(url, download=True)
        filename = ydl.prepare_filename(info)
        print("filename",filename)
    return jsonify({"filename":filename})
# ***********************************FOR PROFRESS BAR******************************************************************
@app.route('/progress')
def get_progress():
    global progress
    with progress_lock:
        return jsonify({'progress': progress})

def progress_hook(d):
    global progress
    if d['status'] == 'downloading':
        with progress_lock:
            progress = d.get('downloaded_bytes', 0) / d.get('total_bytes', 1) * 100


#****************---------NOTES----------*************************************************    

prompt="provide notes of given text "
@app.get("/notes")
def notes():
    return gen_notes(text,prompt)
def gen_notes(txt,prompt):
    global notes
    genai.configure(api_key="AIzaSyB3GLDQavn3ntqNyBsneyyb7N8n9r5Gk1k") #api gemini pro (site=https://aistudio.google.com/app/apikey?_gl=1*7b62cr*_ga*MTc5MTAwMzgyLjE3MjYzODY4MzU.*_ga_P1DBVKWT6V*MTcyNjM4NjgzNS4xLjEuMTcyNjM4Njg3MS4yNC4wLjEwNTUzOTUzMzM. )
    model=genai.GenerativeModel("gemini-pro")
    response=model.generate_content(prompt+txt)
    print(response.text)
    notes= "...................NOTES:-............................." + response.text + "................................................"
    return response.text , 400



# **********************************FILE SAVE FUNCTION****************************************************************

@app.get("/save")
def save():
    global translated_text
    global notes

    # Root window code
    root = tk.Tk()
    root.title("Format Selector")
    root.geometry('300x200')
    root.configure(bg="black")

    style = ttk.Style()
    style.theme_use('clam')

    selected_format = tk.StringVar(root)
    selected_format.set(".text")  # default value

    formats = ['.text', '.pdf', '.docx']  # Formats list

    language_dropdown = ttk.Combobox(root, textvariable=selected_format, values=formats, state="readonly")
    language_dropdown.pack(pady=20, padx=10)

    translate_button = tk.Button(root, text="Submit", command=lambda: [root.destroy()])
    translate_button.pack(pady=10, padx=10)
    translate_button.config(bg='aquamarine', fg='black', font=('times new roman', 10, 'bold'))

    label = tk.Label(root, text="Select a format:", font=('times new roman', 15))
    label.config(bg="violet")
    label.pack(pady=(10, 0))

    root.mainloop()

    # Fetching video title for giving name to file
    video_url = "https://www.youtube.com/watch?v=" + video_id
    yt = YouTube(video_url)

    

    def get_youtube_video_title(video_id):
        api_key="AIzaSyCKLTZ0JiROCVda01VfF2SbuwPNWv-kLVw"
        # YouTube Data API endpoint
        url = f"https://www.googleapis.com/youtube/v3/videos?id={video_id}&key={api_key}&part=snippet"

       
        response = requests.get(url)
        if response.status_code == 200:
            video_data = response.json()
            if video_data['items']:
                title = video_data['items'][0]['snippet']['title']
                return title
            else:
                return "Video not found."
        else:
            return f"Error: {response.status_code}"
    
    v_title= get_youtube_video_title(video_id)
    video_title =  re.sub(r'\W+', '', v_title)
    print(f"Video Title: {video_title}")

    # try:
    #     print(yt.title)
    #     video_title =  re.sub(r'\W+', '', yt.title)
    #     print(f"vid_info: {yt.vid_info}") 
    # except Exception as e:
    #     print(f"An error occurred: {e}")
    #     print(f"vid_info: {yt.vid_info}") 
    #     video_title ="Recent video "
        
   

    # Creating file name
    file_name = ''.join((video_title + selected_format.get()).split())

    # Format conditions
    if selected_format.get() == ".text":
        with open(file_name, "w", encoding="utf-8") as file:
            if text:
                if type(translated_text)!= str and type(notes) != str:
                    file.write(text)
                elif type(translated_text) == str and  type(notes) != str:
                    file.write(text + "\n" + translated_text)
                elif type(translated_text) != str and  type(notes) == str:
                    file.write(text + "\n" + notes)
                elif type(translated_text) == str and type(notes) == str:
                    file.write(text + "\n\n" +translated_text+"\n\n"+ notes)
            else:
                file.close()

    elif selected_format.get() == ".pdf":
        # Save as PDF
        def wrap_text(text, width, font_name, font_size):
            lines = []
            line = ''
            for word in text.split():
                if pdfmetrics.stringWidth(line + ' ' + word, font_name, font_size) <= width:
                    line += ' ' + word
                else:
                    lines.append(line.strip())
                    line = word
            lines.append(line.strip())
            return lines
        font_name="Helvetica"
        if translated_text:
            print("hrllo")
            def get_font_for_language(text):
                language_code = detect(text)
                # selecting language 
                if language_code == 'ta':  # Tamil
                    return 'NotoSansTamil', 'C:/Users/suraj/OneDrive/Desktop/Noto_Sans_Tamil/NotoSansTamil-VariableFont_wdth,wght.ttf'
                elif language_code == 'gu': # gujrati
                    return 'Noto Sans Gujarati Thin','C:/Users/suraj/OneDrive/Desktop/Noto_Sans_Gujarati/NotoSansGujarati-VariableFont_wdth,wght.ttf'
                elif language_code == "bn": # Bengali
                    return 'Noto Sans Bengali Thin','C:/Users/suraj/OneDrive/Desktop/Noto_Sans_Bengali/NotoSansBengali-VariableFont_wdth,wght.ttf'
                elif language_code == 'hi':  # Hindi
                    return 'Tiro Devanagari Hindi',  'C:/Users/suraj/OneDrive/Desktop/Tiro_Devanagari_Hindi/TiroDevanagariHindi-Regular.ttf'
                elif language_code == "mr":# Marathi
                    return 'Tiro Devanagri Marathi','C:/Users/suraj/OneDrive/Desktop/Tiro_Devanagari_Marathi/TiroDevanagariMarathi-Regular.ttf'
                elif language_code == 'te' : #Telugu
                    return 'Noto Serif Telugu Thin','C:/Users/suraj/OneDrive/Desktop/Noto_Serif_Telugu/NotoSerifTelugu-VariableFont_wght.ttf'
                elif language_code == 'zh-cn':  # Chinese
                    return 'Noto Sans TC Thin', "C:/Users/suraj/OneDrive/Desktop/Noto_Sans_SC,Noto_Sans_TC/Noto_Sans_TC/NotoSansTC-VariableFont_wght.ttf"
                else:  # Default  for unsupported languages
                    return 'Helvetica', None

            font_name , font_path =get_font_for_language(translated_text)
            pdfmetrics.registerFont(TTFont(font_name,font_path))
        
        documentTitle = 'YouTube Summerizer'
        title = 'YouTube Summerizer'
        subTitle = video_title
        if text:
            if type(translated_text)!= str and type(notes) != str:
                textLines = [text]
            elif type(translated_text) == str and  type(notes) != str:
                textLines = [text + "\n\n" + translated_text]
                
            elif type(translated_text) != str and  type(notes) == str:
                textLines = [text + "\n\n" + notes]
            elif type(translated_text) == str and type(notes) == str:
                textLines = [text + "\n\n" + translated_text + "\n\n" + notes]
        else:
            textLines="SORRY!!!! \n No Summary "
        width, height = letter
        pdf = canvas.Canvas(file_name, pagesize=letter)
        pdf.setTitle(documentTitle)

        pdf.setFont("Helvetica-Bold", 24)
        pdf.drawCentredString(width / 2, height - 100, title)

        pdf.setFont("Helvetica", 18)
        pdf.drawCentredString(width / 2, height - 150, subTitle)

        pdf.line(30, height - 160, width - 30, height - 160)
        text_y = height - 180
        text_width = width - 60
        font_size = 14
        pdf.setFont(font_name, font_size)

        pdf.setFillColor(colors.black)
        for line in textLines:
            wrapped_lines = wrap_text(line, text_width, font_name, font_size)
            for wrapped_line in wrapped_lines:
                pdf.drawString(40, text_y, wrapped_line)
                text_y -= 20  # Move to the next line
            text_y -= 10  # Add extra space between paragraphs
        pdf.save()

    elif selected_format.get() == ".docx":
        # Save as DOCX
        doc = Document()
        if text:
            if type(translated_text)!= str and type(notes) != str:
                doc.add_paragraph(text)
            elif type(translated_text) == str and  type(notes) != str:
                doc.add_paragraph(text + "\n" + str(translated_text))
            elif type(translated_text) != str and  type(notes) == str:
                doc.add_paragraph(text + "\n" + notes)
            elif type(translated_text) == str and type(notes) == str:
                doc.add_paragraph(text + "\n\n" + str(translated_text) + "\n\n" + notes)
        else:
            doc.add_paragraph("SORRY!!!, No Summary")
        
        doc.save(file_name)

    # Check if file exists and return appropriate response
    file_path = os.path.abspath(file_name)
    if os.path.exists(file_path):
        os.startfile(file_path)
        return send_file(file_path, as_attachment=True), 200
    else:
        return jsonify({"error": "File not found"}), 400
    

    




if __name__ == '__main__':
    app.run()
    